# 创建该文件的LLM大模型：Arena.ai Agent Mode（早期版本）
# 修改该文件的LLM大模型：Claude Sonnet 4.5 (via Arena.ai Agent Mode)
# 最后修改时间（北京时间，精确到秒）：2026-05-30 17:30:00 CST
#
# 修改记录：
# - 2026-05-30 17:30 Claude Sonnet 4.5: 给 §8.4 锚点对应的现状段加上文件头规范新规则的同步内容
"""
把"V1 实际实现现状"子段原地插入到项目需求文档 V1 的 docx。

策略：
1. 保留原 docx 中每段原文一字不改。
2. 在指定的"锚点段落"之后，插入若干段 V1 实际实现现状文本。
3. 重复运行幂等：清理"上一次插入的现状段块"再插入新版。
   - 清理判定：以 STATUS_HEADER_PREFIX 开头算块首；之后连续行只要不与"已知章节标题集合"或"原文已知行"重合，就算块成员，全部清掉。
4. 用于把 md 镜像 (docs/隔音窗专家咨询与购物辅助Agent系统-项目需求文档-V1.md) 中的同步信息回写到 docx。

锚点匹配：以原文中各章标题行的精确字串为锚（见 ANCHOR_BLOCKS）。
触发条件：仅在 ADR-010 列出的 4 种情况（需求/架构/路由/范围变化）下手动跑一次。
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC_PATH = Path(__file__).resolve().parent.parent / "docs" / "隔音窗专家咨询与购物辅助Agent系统-项目需求文档-V1.docx"

# 每个块：(锚点段落精确文本, 在锚点之后追加的段落文本列表)
# 锚点选择策略：每章选一段"内容已结束、马上要进入下一章"的最后一行，作为现状段的插入点。
# 段落文本以 "【V1 实际实现现状（2026-05-30 同步）】" 开头，便于幂等清理。

STATUS_HEADER_PREFIX = "【V1 实际实现现状"

ANCHOR_BLOCKS: List[Tuple[str, List[str]]] = [
    # --- §1.2 核心价值 之后 ---
    (
        "隐私安全：本地化部署，账号Cookie加密隔离，操作可审计",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "🟡 真实比价：淘宝公开页面抓取链路已实现（Playwright + 选择器 + 回退 + 探针 + 反爬骨架）；🚫 拼多多 V1+V1.5 不做（ADR-002 / ADR-009 沿用禁区）。",
            "⏳ 专业诊断 / 个性化推荐 / 避坑指南：Phase 1B 由 ADR-009 承接（A 噪音 / B 材料 / E 施工 + F 协调员），落地后即满足。",
            "✅ 隐私安全：本地优先；URL 白名单 + 危险路径拦截；登录态保存在本地浏览器 profile，不入库。",
            "🔁 与原文 §3.1「Claude Code 作为开发助手」冲突 → 由 ADR-001 替代：不使用 Claude Code，开发由 Arena.ai Agent Mode 主导。",
        ],
    ),
    # --- §2.1 欢迎菜单 之后 ---
    (
        "E. 历史方案管理：浏览、对比、恢复历史咨询记录",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "🚫 Slash Command 入口：V1 不做（无聊天前端，V2+ 才考虑）。",
            "✅ 模式 C「只比价/看商品」：CLI + API + 最小 Web 页面均可触发购物子链路。",
            "⏳ 模式 A「快速评估」/ B「完整方案」/ D「验收指导」：Phase 1B 由 ADR-009 承接，实装后即可对外暴露。",
            "🟡 模式 E「历史方案管理」：购物 run 的历史 / 对比 / 分析已实现（run_compare、run_analysis、phase1_cli history）；会话级历史 Phase 1B 落地后补齐。",
        ],
    ),
    # --- §2.2 核心工作流 之后 ---
    (
        "追问与回溯：点击看板任意卡片可查看详情浮层，或发起追问",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "⏳ F 协调员状态机：Phase 1B 由 ADR-009 承接（src/coordinator/state_machine.py）；当前仅 schema 层定义了咨询上下文。",
            "✅ D 购物参谋：完整购物子链路已实现（intent → keyword → 搜索 → 过滤 → 详情 → 提取 → 排序 → 评论增强 → 二次排序 → LLM 对比总结 → 缓存 → 报告导出）。",
            "🚫 WebSocket 实时推送 / 看板卡片浮层：V2+ 才考虑。",
            "🟡 最小开发态 Web 页面已实装（Jinja2，9 个模板：dashboard / run_detail / run_analysis / compare_runs / tools / artifact_manifest / artifact_detail / event_log / base）。",
        ],
    ),
    # --- §2.3 其他特性 之后 ---
    (
        "购物缓存手动刷新",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "🚫 Markdown 对话渲染 / 暗黑模式：V1 不做（无对话前端）。",
            "⛔ 专家健康监控：V1 不做。",
            "🟡 购物缓存：SQLite 已实现；当前刷新方式是「每次运行写新 run」，未做「针对关键词的手动刷新」按钮。",
        ],
    ),
    # --- §3.1 架构原则 之后 ---
    (
        "轻量隔离：不使用Docker（避免macOS GPU性能损失），改用Python虚拟环境 + 网络代理沙箱 + 文件权限控制",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "✅ FastAPI 后端骨架（src/api/app.py，380 行）。",
            "🔁 Claude Code → 被 ADR-001 替代：开发流程改由 Arena.ai Agent Mode 主导，不使用 Claude Code。",
            "✅ 本地优先；商业模型 deepseek-chat 配置存在但默认 enabled: false。",
            "✅ 不用 Docker；用 uv 管理虚拟环境。",
            "🟡 网络代理沙箱：URL 白名单 + 危险路径拦截已实现（src/security/url_guard.py）；完整本地代理拦截层未实装（当前用直连 + Playwright 上下文隔离 + 选择器约束）。",
        ],
    ),
    # --- §3.3 外部服务 之后 ---
    (
        "购物平台：淘宝、拼多多公开网页（通过安全代理访问）",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "🟡 DeepSeek / GLM：配置存在但默认禁用，触发条件待 OQ-004 决策。",
            "✅ 淘宝（公开页面）。",
            "🚫 拼多多：V1 不做（ADR-002）。",
        ],
    ),
    # --- §4.1 F 协调员 之后 ---
    (
        "健康管理：定期检测各专家服务状态，结果推送看板",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "⏳ 状态机：Phase 1B 由 ADR-009 承接，落地于 src/coordinator/state_machine.py。",
            "✅ 槽位映射：在 src/shopping/intent_builder.py 中以「咨询上下文 → 购物意图」承接；Phase 1B 增加 src/coordinator/slots.py 做交互式追问。",
            "🚫 健康监控：V2+ 才考虑。",
        ],
    ),
    # --- §4.2 A 噪音分析 之后 ---
    (
        "建议隔声量（dB）",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "⏳ A 专家：Phase 1B 由 ADR-009 承接，落地于 src/experts/noise_analyst.py；评测路径 eval_cases/noise_analysis_cases.json 复用。",
            "🔁 路由升级：实际选用 qwen3.6:35b-a3b-q8_0 而非 qwen3:14b（ADR-003）。",
        ],
    ),
    # --- §4.3 B 材料顾问 之后 ---
    (
        "知识源引用附在输出中",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "⏳ B 专家纯 prompt 版：Phase 1B 由 ADR-009 承接，落地于 src/experts/material_advisor.py。",
            "⏳ B 专家 RAG 化：V1.5 由 ADR-012 承接（ChromaDB + 嵌入 + 知识源引用）。",
            "🔁 路由升级：实际由 qwen3.6:35b-a3b-q8_0 承担（ADR-003）。",
        ],
    ),
    # --- §4.4 D 购物参谋 之后 ---
    (
        "缓存策略：手动刷新",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "✅ 核心已实现：搜索 → 列表 → 候选过滤 → 详情 → 字段提取（含回退） → LLM 字段补归纳 → 排序 → LLM 对比总结 → SQLite 缓存 → Markdown 报告。",
            "✅ URL 白名单 / 危险路径拦截 / GET 限制：src/security/url_guard.py。",
            "✅ 评论增强（第二阶段）：review_pipeline + review_enricher + review_fetcher（回放已通，真实抓取骨架已建）。",
            "✅ 反爬：anti_bot_policy + risk_detection（验证码 / 访问受限页面探测）。",
            "🟡 登录态：本地 Playwright user data 目录复用；内存加密 Cookie 未做。",
            "🚫 加购 / 收藏 / 下单 / 付款：V1 暂不做（ADR-002，用户原话）。",
            "🚫 拼多多：V1 不做。",
            "🟡 缓存刷新：每次运行写新 run；未做「针对关键词的手动刷新」按钮。",
            "🔁 路由升级：实际购物推理 / 总结用 qwen3-coder-next:q4_K_M，字段补归纳用 qwen3:14b（ADR-003）。",
        ],
    ),
    # --- §4.5 E 施工指导 之后 ---
    (
        "常见坑点与维护建议",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "⏳ E 专家：Phase 1B 由 ADR-009 承接，落地于 src/experts/installation_guide.py，纯 prompt。",
        ],
    ),
    # --- §4.6 知识库 之后 ---
    (
        "更新策略：初期手动精选资料，使用watchdog监控知识库目录，文件变更后自动重建向量索引（热加载）",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "⏳ 整节移至 V1.5（ADR-012）。",
            "V1.5 范围：ChromaDB（默认）/ FAISS（fallback）；qwen3-embedding:8b 嵌入；手动入库 10~20 篇；六维度评分插件化；watchdog 热加载；B 专家 RAG 化；视频管线最低优先。",
            "仅在 model_router.yaml 保留嵌入模型路由占位。",
        ],
    ),
    # --- §4.7 看板 之后 ---
    (
        "对话区：支持Markdown渲染、代码高亮、图片展示",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "🟡 最小开发态 Web 页面已实装（Jinja2 templates，9 个页面：dashboard / run_detail / run_analysis / compare_runs / tools / artifact_manifest / artifact_detail / event_log / base）。",
            "🚫 Alpine.js / 暗黑模式 / 卡片浮层 / 专家健康状态：V1 不做。",
            "🚫 对话区 / Markdown 渲染：V1 不做（V1 没有聊天前端）。",
        ],
    ),
    # --- §5.1 数据库选型 之后 ---
    (
        "SQLite（本地文件存储），通过 SQLAlchemy ORM 操作，确保可轻松切换其他数据库",
        [
            "【V1 实际实现现状（2026-05-30 同步）— 数据库选型】",
            "✅ SQLite：src/shopping/sqlite_cache.py（232 行）。",
            "🟡 当前未引入 SQLAlchemy ORM，直接用 sqlite3 + dataclass schema（cache_models.py），轻量但牺牲一些可移植性。",
        ],
    ),
    # --- §5.2 核心表结构（在最后一张表"用户偏好表"前的标题行之后插入；这里改为在 §5 末尾，用 6. 安全方案 之前的最后一行原文做锚点）---
    # 选 §6.3 之后的「安全测试用例」头作为锚点
    (
        "6.3 安全测试用例（全维度覆盖）",
        [
            "【V1 实际实现现状（2026-05-30 同步）— 数据模型与安全】",
            "🟡 实际表以「run 为中心」组织（runs / step_traces / events / artifacts / reports），与原 7 表结构不同。",
            "⛔ sessions / messages / solution_versions / knowledge_references / user_preferences：V1 未建表。",
            "✅ URL 白名单 + 危险路径拦截 + 详情页 URL 规范化。",
            "✅ 反爬节流策略（anti_bot_policy + politeness）。",
            "🟡 商业模型 model_guard：未单独抽出模块（因商业模型默认禁用）。",
            "🟡 Cookie 加密 / 日志滚动：V1 未实装。",
            "🟡 安全测试用例 S-01~S-10：仅 S-03 / S-08 / S-09（URL 拦截）有覆盖（test_url_guard.py）。",
        ],
    ),
    # --- §7 模型分配 之后（用最后一行规则作锚）---
    (
        "商业模型调用需记录专用日志",
        [
            "【V1 实际实现现状（2026-05-30 同步）— 模型路由（已被 ADR-003 升级）】",
            "协调员 / 噪音分析 / 方案顾问：主 qwen3.6:35b-a3b-q8_0，fallback qwen3:14b，escalation deepseek-chat。",
            "购物推理 / 购物总结：主 qwen3-coder-next:q4_K_M，fallback qwen3.6:35b-a3b-q8_0，escalation deepseek-chat。",
            "购物字段补归纳：主 qwen3:14b，fallback qwen3-coder-next:q4_K_M。",
            "浏览器执行器：主 deterministic_playwright，fallback headed_playwright_with_manual_takeover。",
            "嵌入：主 qwen3-embedding:8b，fallback bge-m3:latest。",
            "项目开发流程：arena-ai-agent-mode（不使用 Claude Code，ADR-001）。",
            "escalation 触发：完全手动，仅用户主动提出时触发（ADR-011）。CLI / API 将预留 --use-cloud-escalation 开关；实装前必须先做 src/security/model_guard.py 强制剥离敏感信息。",
        ],
    ),
    # --- §8.1 目标环境 之后 ---
    (
        "预装：Python 3.11，Ollama V0.23.1（已拉取 bge-m3, qwen3-embedding:8b, qwen3:14b, qwen3.6:35b-a3b-q8_0），VS Code，Claude Code for VS Code（已可接入任意开源模型）， uv",
        [
            "【V1 实际实现现状（2026-05-30 同步）】",
            "🔁 Claude Code 已被 ADR-001 移除。其余目标环境保持。",
            "✅ uv 已使用（见 pyproject.toml）。",
        ],
    ),
    # --- §8.4 代码质量标准 之后（2026-05-30 第五次新增锚点） ---
    (
        "统一日志：logging模块，同时输出控制台与文件，异常必须记录完整堆栈",
        [
            "【V1 实际实现现状（2026-05-30 第五次同步）— 代码质量标准】",
            "✅ 文件头部注释规范在所有新文件中遵守（见 docs/ONBOARDING_CHECKLIST.md 第 7 步）。",
            "✅ 2026-05-30 第五次升级文件头规范（用户硬性要求）：必须写具体大模型名（如 Claude Sonnet 4.5 (via Arena.ai Agent Mode) / GPT-5 Pro (via Arena.ai Agent Mode) / Gemini 2.5 Pro (via Arena.ai Agent Mode) 等），不能只写笼统的 \"Arena.ai Agent Mode\"，否则无法追溯责任。",
            "✅ 新增文件用：# 创建该文件的LLM大模型：<名> / # 创建时间（北京时间，精确到秒）：YYYY-MM-DD HH:MM:SS CST。",
            "✅ 修改文件用：# 修改该文件的LLM大模型：<名> / # 最后修改时间... 并维护「修改记录」小段（最新在最上）。",
            "✅ 类型注解 / 中文 docstring / 抽象接口（ShoppingExecutorInterface）。",
            "🟡 统一日志：基础 logging 已用，未做完整双输出 + 滚动配置。",
        ],
    ),
    # --- §8.3 技术选型 之后 ---
    (
        "中文亲和：面向用户的输出强制中文；代码注释、docstring用中文；异常日志保留英文原文+中文注释；README中文编写。",
        [
            "【V1 实际实现现状（2026-05-30 同步）— 技术选型】",
            "✅ FastAPI + Jinja2。",
            "⛔ HTMX / Alpine.js / Tailwind：V1 仅最小内联样式。",
            "🟡 SQLite ✅；SQLAlchemy ORM ⛔ 未引入。",
            "✅ uv。",
            "⛔ MLX / MPS 自动回退：未实装（依赖 Ollama 自管）。",
            "✅ 中文亲和（注释 / docstring / 文档全中文；错误日志保留英文 + 中文注释）。",
        ],
    ),
    # --- §8.6 配置管理 之后 ---
    (
        "支持通过 config.yaml 覆盖默认值",
        [
            "【V1 实际实现现状（2026-05-30 同步）— 配置管理】",
            "✅ config.yaml + .env.example + model_router.yaml 三件套。",
            "✅ src/config.py 统一加载。",
            "✅ runtime/selector_overrides.yaml 联调期热覆盖选择器（联调改选择器无需重启）。",
        ],
    ),
    # --- §9 实施路线图 之后（用表 10 的最后一行 Phase 5 作锚）---
    (
        "Phase 5 | 优化增强 | 视频处理管线、自动评分系统、DeepSeek/GLM兜底、用户偏好学习",
        [],  # 表格不便用段落锚点，跳过；改用下面 §9 标题作锚
    ),
    (
        "9. 实施路线图",
        [
            "【V1 实际实现现状（2026-05-30 同步）— 实施路线已重排，ADR-009 + ADR-012】",
            "Phase 0 模型评测与路由定稿：✅ 已完成（两轮评测，结论见 model_router.yaml）。",
            "Phase 1A 淘宝购物主链路 MVP（当前阶段）：🟡 收尾中。联调启动判据见 ADR-008。",
            "Phase 1B 咨询主链路补齐（V1 范围内）：⏳ 待启动，前置依赖 Phase 1A 联调通过。范围：F 协调员状态机 + A 噪音 + B 材料（纯 prompt）+ E 施工 + BaseExpert。",
            "Phase 1C 咨询↔购物 整合闭环：⏳ 待启动，Phase 1B 完成后。",
            "V1.5 知识库 / RAG / 视频管线：⏳ 待启动（ADR-012），前置依赖 V1 全部交付。",
            "V2+ 商业 escalation 自动化 / 偏好学习 / 拼多多 / 看板暗黑 / WebSocket / 健康监控：🚫 V1+V1.5 不做。",
            "Phase 1A 已落地能力（截至 2026-05-30，69 项测试全通过）：意图构建、关键词构建、列表过滤、详情提取（含回退）、LLM 字段补归纳、LLM 对比总结、排序、SQLite 缓存、历史/对比/分析、Markdown 报告、评论增强、反爬骨架、URL 防护、选择器中心化+override、4 种探针、运行级 artifact、handoff 快照、CLI + FastAPI + 最小 Web 页面、Playwright 执行器 MVP（未实机验证）。",
            "Phase 1A 联调前待完成（详见 BACKLOG.md）：API 真实接口统一签名；Markdown 报告统一导出；anti-bot 真正接入执行链；risk_detection 接入中断分支；playwright_executor dry-run 测试；联调清单打磨；scripts/make_patch.py 补丁机制（ADR-013）。",
        ],
    ),
    # --- §10 扩展性设计 之后 ---
    (
        "购物平台搜索工具抽象为ShoppingPlatform接口，新增平台（如京东）只需实现该接口",
        [
            "【V1 实际实现现状（2026-05-30 同步）— 扩展性】",
            "⏳ BaseExpert：Phase 1B 由 ADR-009 承接，落地于 src/experts/base.py。",
            "✅ 购物执行器抽象：src/shopping/executor_interface.py（ShoppingExecutorInterface），已有 playwright_executor / replay_executor 两实现。",
            "⏳ 知识库评分插件化：V1.5 由 ADR-012 承接。",
            "⏳ 槽位配置化：Phase 1B 由 ADR-009 承接，落地于 src/coordinator/slots.py。",
        ],
    ),
]


def find_paragraph_index(doc: Document, target_text: str) -> int:
    """在文档段落中精确匹配 target_text；返回 0 基索引；找不到返回 -1。"""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == target_text.strip():
            return i
    return -1


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """在指定段落后插入新段落，沿用原段落 style。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = paragraph.style
    except Exception:
        # 某些段落 style 拿不到时，保持默认即可
        pass
    new_para.add_run(text)
    return new_para


def _build_known_originals(doc: Document) -> set:
    """收集所有锚点段落的精确文本，作为"原文边界"集合，用于清理时识别下一段原文。"""
    return {anchor for anchor, _ in ANCHOR_BLOCKS}


# 已知原文章节标题行（精确匹配；用于识别"清理时碰到下一节原文"立刻停止）
KNOWN_ORIGINAL_BOUNDARIES = {
    "1. 产品概述", "1.1 产品愿景", "1.2 核心价值",
    "2. 用户场景与功能", "2.1 欢迎菜单（Slash Command /soundproof）",
    "2.2 核心工作流（完整方案模式）", "2.3 其他特性",
    "3. 系统架构", "3.1 架构原则", "3.2 架构全景图", "3.3 外部服务",
    "4. 功能模块详述", "4.1 F 协调员", "4.2 A 噪音分析专家",
    "4.3 B 材料与方案顾问", "4.4 D 预算与购物参谋",
    "4.5 E 施工与验收指导", "4.6 知识库模块", "4.7 看板与交互",
    "5. 数据模型", "5.1 数据库选型", "5.2 核心表结构",
    "6. 安全方案", "6.1 安全层级", "6.2 商业模型使用红线",
    "6.3 安全测试用例（全维度覆盖）",
    "7. 模型分配与降级策略",
    "8. 开发与工程规范", "8.1 目标环境", "8.2 项目结构",
    "8.3 技术选型约束", "8.4 代码质量标准",
    "8.5 中国网络环境适配", "8.6 配置管理",
    "9. 实施路线图",
    "10. 扩展性设计",
    "11. 另外一个技术栈组合的想法：",
    "#我有多认真：", "#项目框架：", "#原则：",
}


def remove_existing_status_block(doc: Document, anchor_idx: int) -> int:
    """从 anchor_idx+1 起，删掉连续的"V1 实际实现现状"段落块。

    停止条件（满足其一即停止）：
    - 文档末尾；
    - 当前行匹配某个已知原文章节标题（KNOWN_ORIGINAL_BOUNDARIES）；
    - 当前行匹配另一个 anchor（说明已进入下一节内容）；
    - 连续遇到 2 个空行（保守边界）。

    第一段要求以 STATUS_HEADER_PREFIX 开头，否则什么都不删。
    """
    other_anchors = {a for a, _ in ANCHOR_BLOCKS}
    removed = 0
    consecutive_blank = 0
    while True:
        target_idx = anchor_idx + 1
        if target_idx >= len(doc.paragraphs):
            break
        p = doc.paragraphs[target_idx]
        text = p.text.strip()
        if removed == 0:
            # 第一段必须是现状段块头，否则不动
            if not text.startswith(STATUS_HEADER_PREFIX):
                break
        else:
            if text in KNOWN_ORIGINAL_BOUNDARIES:
                break
            if text in other_anchors and text != "":
                break
            if text == "":
                consecutive_blank += 1
                if consecutive_blank >= 2:
                    break
            else:
                consecutive_blank = 0
        p._p.getparent().remove(p._p)
        removed += 1
    return removed


def apply_sync(doc_path: Path) -> None:
    doc = Document(str(doc_path))

    inserted_blocks = 0
    skipped_blocks = 0

    for anchor_text, status_lines in ANCHOR_BLOCKS:
        if not status_lines:
            continue
        anchor_idx = find_paragraph_index(doc, anchor_text)
        if anchor_idx == -1:
            print(f"  [WARN] 锚点未找到，跳过：{anchor_text[:50]}...")
            skipped_blocks += 1
            continue

        # 幂等：先清理同位置已存在的现状段块
        removed = remove_existing_status_block(doc, anchor_idx)
        if removed > 0:
            print(f"  [INFO] 清理旧现状段 {removed} 行 @ 锚点「{anchor_text[:30]}...」")

        # 重新定位 anchor（删除后 index 可能变）
        anchor_idx = find_paragraph_index(doc, anchor_text)
        if anchor_idx == -1:
            print(f"  [ERROR] 清理后锚点丢失：{anchor_text[:30]}")
            continue

        anchor_para = doc.paragraphs[anchor_idx]

        # 倒序插入（addnext 是紧跟在 anchor 之后），所以从最后一行开始反向插，最终顺序正确
        for line in reversed(status_lines):
            insert_paragraph_after(anchor_para, line)
        inserted_blocks += 1
        print(f"  [OK] 已插入 {len(status_lines)} 行 @ 锚点「{anchor_text[:30]}...」")

    doc.save(str(doc_path))
    print(f"\n完成：插入 {inserted_blocks} 个块，跳过 {skipped_blocks} 个块。")
    print(f"已保存到：{doc_path}")


if __name__ == "__main__":
    if not DOC_PATH.exists():
        print(f"找不到 docx：{DOC_PATH}", file=sys.stderr)
        sys.exit(1)
    apply_sync(DOC_PATH)
